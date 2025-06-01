import base64
import os
import json
from pathlib import Path

from azure.search.documents._generated.models import VectorizableTextQuery, VectorizedQuery
from azure.storage.blob import BlobServiceClient
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexerClient
from azure.core.credentials import AzureKeyCredential
import time

import openai
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

load_dotenv()
AZURE_STORAGE_CONN            = os.getenv("AZURE_STORAGE_CONN")
AZURE_SEARCH_ENDPOINT         = os.getenv("SEARCH_ENDPOINT")
AZURE_SEARCH_KEY              = os.getenv("SEARCH_KEY")
REGULATION_SEARCH_IDX_NAME    = os.getenv("REGULATION_SEARCH_IDX_NAME")
REPORT_SEARCH_IDX_NAME        = os.getenv("REPORT_SEARCH_IDX_NAME")
REPORT_INDEXER_NAME           = os.getenv("REPORT_INDEXER_NAME")

EMBEDDING_DEPLOYMENT   = os.getenv("EMBED_DEPLOY")
CHAT_DEPLOYMENT        = os.getenv("CHAT_DEPLOY")

AZURE_OPENAI_ENDPOINT  = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY   = os.getenv("AZURE_OPENAI_KEY")

# Azure Blob Storage client
blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONN)
report_container     = blob_service_client.get_container_client("reports")
reg_container        = blob_service_client.get_container_client("raw")

# Azure Search clients
search_credential    = AzureKeyCredential(AZURE_SEARCH_KEY)
search_client        = SearchClient(
    endpoint=AZURE_SEARCH_ENDPOINT,
    index_name=REGULATION_SEARCH_IDX_NAME,
    credential=search_credential
)
indexer_client       = SearchIndexerClient(
    endpoint=AZURE_SEARCH_ENDPOINT,
    credential=search_credential
)

# Azure OpenAI (via OpenAI Python SDK)
openai.api_key       = AZURE_OPENAI_API_KEY
openai.api_base      = AZURE_OPENAI_ENDPOINT
openai.api_type      = "azure"
openai.api_version   = os.getenv("AZURE_OPENAI_API_VERSION")

ESG_PROMPT_TEMPLATE = """
<instructions>
    Olet ESG‐tarkastuksen asiantuntija. Vertaa seuraavaa otetta yrityksen ESG‐raportista
    annettuun sääntelytekstiin. Molemmat katkelmat kuuluvat samaan pykälään {section}.
    Tunnista:
      1. Epäselvyydet siinä, miten raportti käsittelee sääntelyä.
      2. Kaikki kohdat, jotka ovat ristiriidassa sääntelyn kanssa tai puuttuvat verrattuna sääntelytekstiin.
      3. Suositukset siitä, miten raportti voidaan yhdenmukaistaa sääntelyn kanssa.

    <strong>
        Lisäksi jokaiseen “ambiguity”‐ ja “inconsistency”‐kohtaan lisätään lähdeviite
        muodossa [Lähde: {section}] heti sen perään AINOASTAAN JOS kyseisessä kohdassa ollaan käytetty kyseistä lähdettä.
    </strong>
</instructions>

<report>
    {report_text}
</report>

<regulation>
    {regulation_text}
</regulation>

<strict>
    <response‐format>
        {{
            "ambiguities": [ 
                /* Jokainen merkkijono päättyy lähdeviitteellä: "... [Lähde: {section}]" */
            ],
            "inconsistencies": [
                /* Sama: "... [Lähde: {section}]" */
            ],
            "recommendations": [ ... list of strings ... ]
        }}
    </response‐format>
</strict>
"""


def upload_report_to_blob(local_pdf_path: str):
    """
    Uploads the given local PDF to the 'reports' container.
    Tags it with sourceType='report' and originId equal to the blob name (without extension).
    Returns the blob name and origin_id.
    """
    filename   = os.path.basename(local_pdf_path)
    origin_id  = os.path.splitext(filename)[0]
    blob_client = report_container.get_blob_client(blob=filename)

    with open(local_pdf_path, "rb") as f:
        blob_client.upload_blob(f, overwrite=True)

    try:
        blob_client.set_blob_tags({
            "sourceType": "report",
            "originId":   origin_id
        })
    except Exception:
        pass

    return filename, origin_id

def run_report_indexer(indexer_name: str, poll_interval: int = 5):
    """
    Triggers the given indexer and polls until completion.
    """
    indexer_client.run_indexer(indexer_name)
    print(f"Indeksointi aloitettu indeksoijalle '{indexer_name}'. Haetaan statustietoja...")
    time.sleep(poll_interval)

    while True:
        status = indexer_client.get_indexer_status(indexer_name)
        print(f"  Indeksoijan status: {status.status}")
        if status.last_result.status.lower() in ["running", "inprogress", "in_progress"]:
            time.sleep(poll_interval)
            continue
        if status.status.lower() == "error":
            failures = status.last_result.errors
            raise RuntimeError(f"Indeksoija '{indexer_name}' ei voinut indeksoida: {failures}")
        print(f"Indeksoijan'{indexer_name}' indeksointi valmis.")
        break


esg_search_client = SearchClient(
    endpoint=AZURE_SEARCH_ENDPOINT,
    index_name=REPORT_SEARCH_IDX_NAME,
    credential=AzureKeyCredential(AZURE_SEARCH_KEY)
)

def get_esg_passages(origin_id: str, top_n: int = 5):
    """
    Fetches the top-N ESG‐related passages from the report index,
    identified by origin_id. Expects each chunk to carry a 'section' field.
    """

    url = f"https://{os.getenv('BLOB_STORAGE_NAME')}.blob.core.windows.net/reports/{origin_id}.pdf"
    b64_str = base64.b64encode(bytes(url, 'ascii')).decode('ascii')
    print(b64_str)

    # Build a vector query for ESG‐topics
    vq = VectorizableTextQuery(
        text="Find sections that discuss environmental metrics, social impact, or governance disclosures.",
        fields="text_vector",
        k_nearest_neighbors=top_n
    )
    results = esg_search_client.search(
        search_text="*",
        vector_queries=[ vq ],
        query_type="semantic",
        filter=f"parent_id eq '{b64_str}0'",
        semantic_configuration_name="reports-semantic-configuration"
    )
    passages = []
    for hit in results:
        passages.append({
            "chunk_id": hit["chunk_id"],
            "content":  hit["chunk"],
            "section":  hit.get("section", "")   # Now expects 'section' from index
        })
    return passages

def get_regulation_chunks_by_vector(report_chunk_embedding, top_k: int = 3):
    """
    Performs a vector KNN search against the regulation index using
    the given embedding. Returns the top_k regulation chunks, each
    expected to include a 'section' field.
    """
    vq = VectorizedQuery(
        fields="text_vector",
        vector=report_chunk_embedding,
        k_nearest_neighbors=top_k
    )
    response = search_client.search(
        search_text="*",
        vector_queries=[ vq ],
        top=top_k
    )

    reg_chunks = []
    for hit in response:
        reg_chunks.append({
            "reg_chunk_id": hit.get("chunk_id", hit.get("id")),
            "content":      hit.get("chunk", hit.get("content")),
            "title":      hit.get("title")
        })
    return reg_chunks


def compare_report_to_regulation(report_text: str, regulation_text: str, regulation_source: str):
    """
    Calls the Azure OpenAI Chat endpoint to compare the given report excerpt
    to the regulation excerpt, passing the common `section` string as source.
    Returns the raw model output (string).
    """
    prompt = ESG_PROMPT_TEMPLATE.format(
        report_text=report_text,
        regulation_text=regulation_text,
        section=regulation_source
    )

    response = openai.chat.completions.create(
        model=CHAT_DEPLOYMENT,
        messages=[
            {"role": "system", "content": "Olet ESG‐tarkastusten asiantuntija."},
            {"role": "user",   "content": prompt}
        ],
        max_completion_tokens=4096,
    )
    return response.choices[0].message.content


def pretty_print_results(results_list):
    """
    Pretty‐prints ESG vs. regulation comparison results to the console.
    Each item in results_list is a dict with:
      - 'report_chunk_id'
      - 'reg_chunk_id'
      - 'section'
      - 'analysis': either a 'raw_output' string or structured lists under keys
        like 'ambiguities', 'inconsistencies', 'recommendations'.
    """
    for idx, item in enumerate(results_list, start=1):
        print(f"\n=== Comparison {idx} ===")
        print(f"Report Chunk ID: {item.get('report_chunk_id')}")
        print(f"Regulation Chunk ID: {item.get('reg_chunk_id')}")
        print(f"Section: {item.get('section')}\n")
        analysis = item.get("analysis", {})

        raw = analysis.get("raw_output", "")
        if raw:
            print("Raw Output:")
            try:
                parsed = json.loads(raw)
                print(json.dumps(parsed, indent=4))
            except (json.JSONDecodeError, TypeError):
                print(raw)
        else:
            for key in ("ambiguities", "inconsistencies", "recommendations"):
                values = analysis.get(key)
                if isinstance(values, list) and values:
                    print(f"\n{key.capitalize()}:")
                    for entry in values:
                        print(f"  - {entry}")


def save_comparisons_to_docx(all_comparisons, output_folder="output", source_filename="NA"):
    """
    Creates a .docx file containing only those comparison results with fully parsed JSON.
    Any comparison with an unparsed 'raw_output' is written to 'unparsed.log' instead.
    Each comparison heading uses the 'section' field as the title.
    """
    Path(output_folder).mkdir(parents=True, exist_ok=True)
    timestamp = int(time.time())
    log_path = Path(output_folder) / f"{source_filename}_{timestamp}.log"
    log_file = open(log_path, "w", encoding="utf-8")

    doc = Document()
    doc.add_heading(f"FinReg Agentti – Analyysi raportille \"{source_filename}.pdf\"", level=1)

    for idx, item in enumerate(all_comparisons, start=1):
        report_id = item.get("report_chunk_id", "N/A")
        reg_id    = item.get("reg_chunk_id",    "N/A")
        #title     = item.get("reg_title",         f"Comparison {idx}")
        analysis  = item.get("analysis",        {})
        report_content = item.get("report_content")

        raw = analysis.get("raw_output")
        if isinstance(raw, str) and raw.strip():
            log_file.write(f"Vertailu {idx} - Raporttilohkon ID: {report_id}, Sääntelylohkon ID: {reg_id}")
            log_file.write(raw + "\n\n")
            continue

        # Use 'section' as the heading for each comparison
        #doc.add_heading(title, level=2)
        doc.add_heading(f"Vertailu #{idx}", level=2)
        doc.add_paragraph("Raportin kohta, jota analysoidaan", style="Intense Quote")
        doc.add_paragraph( report_content )

        for key in ("ambiguities", "inconsistencies", "recommendations"):
            values = analysis.get(key)
            if isinstance(values, list) and values:
                heading = ""
                if key == "ambiguities": heading = "Epäselvyydet"
                if key == "inconsistencies": heading = "Ristiriitaisuudet"
                if key == "recommendations": heading = "Suositukset"

                doc.add_paragraph(f"{heading}:", style="Intense Quote")
                for entry in values:
                    bullet = doc.add_paragraph(style="List Bullet")
                    run = bullet.add_run(entry)
                    run.font.size = Pt(10)

        doc.add_page_break()

    log_file.close()


    output_path = Path(output_folder) / f"{source_filename}_{timestamp}.docx"
    doc.save(str(output_path.resolve()))
    print(f"Saved valid comparisons to '{output_path.resolve()}'")
    print(f"Saved unparsed entries to '{log_path.resolve()}'")


if __name__ == "__main__":
    local_pdf_path = ""
    while local_pdf_path == "":
        local_pdf_path = input("ESG-raportin polku: ")

    blob_name, origin_id = upload_report_to_blob(local_pdf_path)
    print(f"Lataus onnistui: '{blob_name}' originId:nä '{origin_id}' Blob Storageen.")

    run_report_indexer(indexer_name=REPORT_INDEXER_NAME)

    report_depth = input("Kuinka monta relevantteja ESG-liitännäisiä lohkoja haluat (oletus: 3): ")
    if not report_depth:
        report_depth = 3
    else:
        report_depth = int(report_depth)

    print(f"Haetaan raportista {report_depth} relevanteinta ESG‐aiheisia lohkoja…")
    esg_passages = get_esg_passages(origin_id, top_n=report_depth)
    if not esg_passages:
        print("ESG‐aiheisia osuuksia ei löytynyt. Tarkista indeksointi ja tagien konfiguraatio.")
        exit(1)

    regulation_depth = input("Kuinka monta relevantteja sääntelylohkoja haluat (oletus: 3): ")
    if not regulation_depth:
        regulation_depth = 3
    else:
        regulation_depth = int(regulation_depth)

    all_comparisons = []
    for passage in esg_passages:
        emb_resp = openai.embeddings.create(
            model=EMBEDDING_DEPLOYMENT,
            input=[passage["content"]]
        )
        rpt_embedding = emb_resp.data[0].embedding

        reg_chunks = get_regulation_chunks_by_vector(rpt_embedding, top_k=regulation_depth)
        if not reg_chunks:
            print(f"Sääntelylohkoja ei löytynyt raportin lohkolle '{passage['chunk_id']}'. Ohitetaan.")
            continue

        for reg in reg_chunks:
            print(f"Vertailen raportin lohkoa '{passage['chunk_id']}' ja sääntelylohkoa '{reg['reg_chunk_id']}'...")
            ai_output = compare_report_to_regulation(
                passage["content"],
                reg["content"],
                reg.get("title", "")
            )
            try:
                parsed = json.loads(ai_output)
            except json.JSONDecodeError:
                parsed = {"raw_output": ai_output}

            all_comparisons.append({
                "report_chunk_id": passage["chunk_id"],
                "reg_chunk_id":    reg["reg_chunk_id"],
                "reg_title": reg["title"],
                "report_content":  passage["content"],
                "analysis":        parsed,
                "section":         reg.get("section", "")
            })

    #print("\n=== Comparison Results (Pretty printed) ===")
    #pretty_print_results(all_comparisons)
    save_comparisons_to_docx(
        all_comparisons,
        source_filename=f"{origin_id}")